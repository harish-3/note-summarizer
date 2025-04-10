"""
File Parser Module for Note Summarizer Application

This module handles the parsing of PDF and DOCX files to extract text content.
"""

import PyPDF2
import docx
import os
from typing import Dict, Any, Optional


class FileParser:
    """Class to handle parsing of different file formats."""
    
    @staticmethod
    def parse_file(file_path: str) -> Dict[str, Any]:
        """
        Parse a file and extract its text content.
        
        Args:
            file_path: Path to the file to be parsed
            
        Returns:
            Dictionary containing the extracted text and metadata
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return FileParser.parse_pdf(file_path)
        elif file_extension == '.docx':
            return FileParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """
        Parse a PDF file and extract its text content.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing the extracted text and metadata
        """
        text_content = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata if available
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    }
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n\n"
        
        except Exception as e:
            raise Exception(f"Error parsing PDF file: {str(e)}")
        
        return {
            'text': text_content.strip(),
            'metadata': metadata,
            'pages': len(pdf_reader.pages) if 'pdf_reader' in locals() else 0,
            'format': 'pdf'
        }
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract its text content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing the extracted text and metadata
        """
        text_content = ""
        metadata = {}
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata if available
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'subject': core_properties.subject or '',
                'comments': core_properties.comments or ''
            }
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
                text_content += "\n"
        
        except Exception as e:
            raise Exception(f"Error parsing DOCX file: {str(e)}")
        
        return {
            'text': text_content.strip(),
            'metadata': metadata,
            'paragraphs': len(doc.paragraphs) if 'doc' in locals() else 0,
            'format': 'docx'
        }


def get_file_extension(filename: str) -> str:
    """
    Get the extension of a file.
    
    Args:
        filename: Name of the file
        
    Returns:
        File extension (lowercase)
    """
    return os.path.splitext(filename)[1].lower()


def is_supported_file(filename: str) -> bool:
    """
    Check if a file is supported by the parser.
    
    Args:
        filename: Name of the file
        
    Returns:
        True if the file is supported, False otherwise
    """
    supported_extensions = ['.pdf', '.docx']
    return get_file_extension(filename) in supported_extensions
